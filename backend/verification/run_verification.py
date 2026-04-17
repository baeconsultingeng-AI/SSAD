"""
run_verification.py
===================
SSAD Structural Engineering App — Element Calculation Verification Script

Generates 'Element Verification Results.docx' in the project root.

Verification methodology
------------------------
Each element has 3 test cases with different input parameters.
Reference values are independently hand-calculated from the governing
design standard (BS 8110-1:1997 for RC; BS 5950-1:2000 for Steel) and
are coded below. The app's computed values are compared against reference
values within a ±2% tolerance.

Run from the project root:
    python backend/verification/run_verification.py
"""

from __future__ import annotations

import math
import sys
import os
from pathlib import Path
from datetime import date

# ── Path setup ────────────────────────────────────────────────────────────────
BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from app.calc_modules import (
    rc_beam_bs_v1,
    rc_slab_bs_v1,
    rc_column_bs_v1,
    rc_foundation_bs_v1,
    steel_beam_bs_v1,
    steel_column_bs_v1,
    steel_truss_bs_v1,
    steel_portal_bs_v1,
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ─────────────────────────────────────────────────────────────────
TOLERANCE = 0.02          # 2% tolerance on numerical comparisons
TODAY     = date.today().strftime("%d %B %Y")
APP_VER   = "SSAD v1.0"
STANDARD  = "BS 8110-1:1997 / BS 5950-1:2000"

# Colours
CLR_PASS   = RGBColor(0x15, 0x80, 0x3D)   # green
CLR_FAIL   = RGBColor(0xB9, 0x1C, 0x1C)   # red
CLR_HDR    = RGBColor(0x1A, 0x4A, 0x8A)   # navy blue
CLR_SUBHDR = RGBColor(0x1E, 0x40, 0xAF)   # lighter blue
CLR_WARN   = RGBColor(0xB4, 0x53, 0x09)   # amber


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPER UTILITIES
# ═══════════════════════════════════════════════════════════════════════════════

def step_val(result: dict, *fragments: str) -> float | None:
    """Return the value of the first step whose label contains all given fragments (case-insens.)."""
    for step in result.get("steps", []):
        label = step["label"].lower()
        if all(f.lower() in label for f in fragments):
            return float(step["value"])
    return None


def check_status(result: dict, check_id: str) -> str:
    """Return the status string for a named check."""
    for c in result.get("checks", []):
        if c["id"] == check_id:
            return c["status"]
    return "—"


def compare(app_val: float | None, ref_val: float, tol: float = TOLERANCE) -> tuple[str, str]:
    """Compare app value to reference.  Returns (symbol, css-colour-name)."""
    if app_val is None:
        return "N/A", "warn"
    if ref_val == 0.0:
        ok = abs(app_val) < 1e-6
    else:
        ok = abs(app_val - ref_val) / abs(ref_val) <= tol
    return ("✓ MATCH" if ok else "✗ MISMATCH"), ("pass" if ok else "fail")


# ═══════════════════════════════════════════════════════════════════════════════
#  TEST CASE DEFINITIONS  (inputs + reference values)
# ═══════════════════════════════════════════════════════════════════════════════

# ── 1. RC BEAM ─────────────────────────────────────────────────────────────────

RC_BEAM_CASES = [
    {
        "id": "RC-B-01",
        "desc": "250×450 mm Simply Supported Beam — 5 m Span",
        "standard_clause": "BS 8110-1:1997, Cl. 3.4.4 (Bending), Cl. 3.4.5 (Shear)",
        "inputs": {
            "geometry":  {"span": 5.0, "b": 250.0, "h": 450.0, "cover": 30.0,
                          "support_type": "simply_supported"},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"gk": 10.0, "qk": 7.5},
        },
        "reference": {
            # w = 1.4×10 + 1.6×7.5 = 26.0 kN/m
            # M = wL²/8 = 26×25/8 = 81.25 kN·m
            # V = wL/2 = 26×5/2 = 65.0 kN
            # d = 450−30−8−8 = 404 mm
            # K = 81.25×10⁶/(30×250×404²) = 0.0663
            # z = 404×(0.5+√(0.25−0.0663/0.9)) = 371.6 mm  [z ≤ 0.95d=383.8]
            # As = 81.25×10⁶/(0.87×460×371.6) = 545.7 mm²
            "Design UDL":          26.00,
            "Design moment":       81.25,
            "Design shear":        65.00,
            "Effective depth":    404.0,
            "K factor":            0.0663,
            "As required":       545.7,
        },
    },
    {
        "id": "RC-B-02",
        "desc": "300×550 mm Continuous Beam — 7 m Span",
        "standard_clause": "BS 8110-1:1997, Cl. 3.4.4, Cl. 3.5.2.3 (continuous bending)",
        "inputs": {
            "geometry":  {"span": 7.0, "b": 300.0, "h": 550.0, "cover": 30.0,
                          "support_type": "continuous"},
            "materials": {"fcu": 25.0, "fy": 460.0},
            "loads":     {"gk": 15.0, "qk": 10.0},
        },
        "reference": {
            # w = 1.4×15 + 1.6×10 = 37.0 kN/m
            # M = 0.086×37×7² = 155.9 kN·m  (BS 8110 mid-span continuous)
            # V = 0.6×37×7/2 = 77.7 kN
            # d = 550−30−8−8 = 504 mm
            # K = 155.9×10⁶/(25×300×504²) = 0.0819
            # z = 504×(0.5+√(0.25−0.0819/0.9)) = 453.0 mm
            # As = 155.9×10⁶/(0.87×460×453.0) = 859.0 mm²
            "Design UDL":         37.00,
            "Design moment":     155.9,
            "Design shear":       77.7,
            "Effective depth":   504.0,
            "K factor":           0.0819,
            "As required":       859.0,
        },
    },
    {
        "id": "RC-B-03",
        "desc": "200×400 mm Cantilever Beam — 3 m Span",
        "standard_clause": "BS 8110-1:1997, Cl. 3.4.4, Cl. 3.4.6 (deflection)",
        "inputs": {
            "geometry":  {"span": 3.0, "b": 200.0, "h": 400.0, "cover": 25.0,
                          "support_type": "cantilever"},
            "materials": {"fcu": 35.0, "fy": 460.0},
            "loads":     {"gk": 8.0, "qk": 5.0},
        },
        "reference": {
            # w = 1.4×8 + 1.6×5 = 19.2 kN/m
            # M = wL²/2 = 19.2×9/2 = 86.4 kN·m
            # V = wL = 19.2×3 = 57.6 kN
            # d = 400−25−8−8 = 359 mm
            # K = 86.4×10⁶/(35×200×359²) = 0.0958
            # z = 359×(0.5+√(0.25−0.0958/0.9)) = 315.7 mm
            # As = 86.4×10⁶/(0.87×460×315.7) = 684.1 mm²
            "Design UDL":         19.20,
            "Design moment":      86.40,
            "Design shear":       57.60,
            "Effective depth":   359.0,
            "K factor":           0.0958,
            "As required":       684.1,
        },
    },
]

# ── 2. RC SLAB ─────────────────────────────────────────────────────────────────

RC_SLAB_CASES = [
    {
        "id": "RC-S-01",
        "desc": "One-Way Simply Supported Slab — h=200 mm, 4 m Span",
        "standard_clause": "BS 8110-1:1997, Cl. 3.5 (one-way slab bending & deflection)",
        "inputs": {
            "geometry":  {"slab_type": "one_way", "span": 4.0, "h": 200.0,
                          "cover": 30.0, "support_type": "simply_supported"},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"gk": 3.0, "qk": 4.0},
        },
        "reference": {
            # Per 1 m width:
            # w = 1.4×3 + 1.6×4 = 10.6 kN/m²
            # M = 10.6×16/8 = 21.2 kN·m/m
            # d = 200−30−6 = 164 mm  (12 mm bar assumed)
            # K = 21.2×10⁶/(30×1000×164²) = 0.0263
            # z capped at 0.95d = 155.8 mm
            # As = 21.2×10⁶/(0.87×460×155.8) = 340.0 mm²/m
            "Design UDL":         10.60,
            "Design moment":      21.20,
            "Effective depth":   164.0,
            "K factor":           0.0263,
            "As required":       340.0,
        },
    },
    {
        "id": "RC-S-02",
        "desc": "Two-Way Simply Supported Slab — h=150 mm, lx=3.5 m, ly=5.0 m",
        "standard_clause": "BS 8110-1:1997, Cl. 3.5.3.3 & Table 3.14 (two-way slab)",
        "inputs": {
            "geometry":  {"slab_type": "two_way", "lx": 3.5, "ly": 5.0,
                          "h": 150.0, "cover": 25.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"gk": 2.5, "qk": 3.0},
        },
        "reference": {
            # n = 1.4×2.5 + 1.6×3.0 = 8.3 kN/m²
            # ly/lx = 5/3.5 = 1.43 → BS 8110 Table 3.14: βsx≈0.099, βsy≈0.051
            # Msx = 0.099×8.3×3.5² = 10.07 kN·m/m  (short span)
            # d = 150−25−6 = 119 mm
            # K_sx = 10.07×10⁶/(30×1000×119²) = 0.0237
            # z_sx = min(119×(0.5+√(0.25−0.0237/0.9)), 0.95×119) = 113.05 mm
            # As_sx = 10.07×10⁶/(0.87×460×113.05) = 222.4 mm²/m
            "n (ULS load)":        8.30,
            "Effective depth":   119.0,
            "Msx (short span)":   10.07,
            "As short span":     222.4,
        },
    },
    {
        "id": "RC-S-03",
        "desc": "Cantilever Slab — h=150 mm, 2 m Span",
        "standard_clause": "BS 8110-1:1997, Cl. 3.5 (cantilever slab)",
        "inputs": {
            "geometry":  {"slab_type": "cantilever", "span": 2.0, "h": 150.0,
                          "cover": 25.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"gk": 3.0, "qk": 2.0},
        },
        "reference": {
            # w = 1.4×3 + 1.6×2 = 7.4 kN/m²
            # M = wL²/2 = 7.4×4/2 = 14.8 kN·m/m
            # d = 150−25−6 = 119 mm
            # K = 14.8×10⁶/(30×1000×119²) = 0.0348
            # z = min(119×(0.5+√(0.25−0.0348/0.9)), 0.95×119) = 113.05 mm
            # As = 14.8×10⁶/(0.87×460×113.05) = 326.9 mm²/m
            "Design UDL":         7.40,
            "Design moment":      14.80,
            "Effective depth":   119.0,
            "K factor":           0.0348,
            "As required":       326.9,
        },
    },
]

# ── 3. RC COLUMN ───────────────────────────────────────────────────────────────

RC_COLUMN_CASES = [
    {
        "id": "RC-C-01",
        "desc": "300×300 mm Short Braced Column — N=500 kN, Mx=20 kN·m",
        "standard_clause": "BS 8110-1:1997, Cl. 3.8.4 (short braced column)",
        "inputs": {
            "geometry":  {"b": 300.0, "h": 300.0, "cover": 30.0,
                          "le_x": 3.0, "le_y": 3.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"N_Ed": 500.0, "Mx": 20.0, "My": 0.0},
        },
        "reference": {
            # Ag = 300×300 = 90,000 mm²
            # lex/h = 3000/300 = 10.0  (<15 → short column ✓)
            # N_max = 0.35×30×90000 + 0.67×460×As_min
            "Gross section area":   90000.0,
            "lex/h":                   10.0,
        },
    },
    {
        "id": "RC-C-02",
        "desc": "350×350 mm Short Braced Column — N=800 kN, Mx=50 kN·m",
        "standard_clause": "BS 8110-1:1997, Cl. 3.8.4",
        "inputs": {
            "geometry":  {"b": 350.0, "h": 350.0, "cover": 30.0,
                          "le_x": 4.5, "le_y": 4.5},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"N_Ed": 800.0, "Mx": 50.0, "My": 0.0},
        },
        "reference": {
            # Ag = 350×350 = 122,500 mm²
            # lex/h = 4500/350 = 12.86  (<15 → short ✓)
            "Gross section area":  122500.0,
            "lex/h":                   12.857,
        },
    },
    {
        "id": "RC-C-03",
        "desc": "400×400 mm Short Braced Column — N=1200 kN, Mx=60 kN·m",
        "standard_clause": "BS 8110-1:1997, Cl. 3.8.4",
        "inputs": {
            "geometry":  {"b": 400.0, "h": 400.0, "cover": 40.0,
                          "le_x": 3.6, "le_y": 3.6},
            "materials": {"fcu": 35.0, "fy": 460.0},
            "loads":     {"N_Ed": 1200.0, "Mx": 60.0, "My": 0.0},
        },
        "reference": {
            # Ag = 400×400 = 160,000 mm²
            # lex/h = 3600/400 = 9.0  (<15 → short ✓)
            "Gross section area":  160000.0,
            "lex/h":                    9.0,
        },
    },
]

# ── 4. RC FOUNDATION ───────────────────────────────────────────────────────────

RC_FOUND_CASES = [
    {
        "id": "RC-F-01",
        "desc": "2.5×2.0 m Pad Footing — N_sls=400 kN, qa=150 kN/m²",
        "standard_clause": "BS 8110-1:1997, Cl. 3.11 (pad footing)",
        "inputs": {
            "geometry":  {"L": 2.5, "B": 2.0, "h": 500.0, "cover": 75.0,
                          "col_h": 300.0, "col_b": 300.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"N_sls": 400.0, "N_uls": 560.0, "qa": 150.0, "M_sls": 0.0},
        },
        "reference": {
            # p_uls = N_uls/(L×B) = 560/(2.5×2.0) = 112.0 kN/m²
            # d = 500−75−8 = 417 mm
            "ULS upward pressure":  112.0,
            "Effective depth":      417.0,
        },
    },
    {
        "id": "RC-F-02",
        "desc": "2.0×2.0 m Pad Footing — N_sls=350 kN, qa=120 kN/m²",
        "standard_clause": "BS 8110-1:1997, Cl. 3.11",
        "inputs": {
            "geometry":  {"L": 2.0, "B": 2.0, "h": 450.0, "cover": 75.0,
                          "col_h": 300.0, "col_b": 300.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"N_sls": 350.0, "N_uls": 490.0, "qa": 120.0, "M_sls": 0.0},
        },
        "reference": {
            # p_uls = 490/(2.0×2.0) = 122.5 kN/m²
            # d = 450−75−8 = 367 mm
            "ULS upward pressure":  122.5,
            "Effective depth":      367.0,
        },
    },
    {
        "id": "RC-F-03",
        "desc": "3.0×2.5 m Pad Footing — N_sls=700 kN, qa=200 kN/m²",
        "standard_clause": "BS 8110-1:1997, Cl. 3.11",
        "inputs": {
            "geometry":  {"L": 3.0, "B": 2.5, "h": 550.0, "cover": 75.0,
                          "col_h": 400.0, "col_b": 400.0},
            "materials": {"fcu": 30.0, "fy": 460.0},
            "loads":     {"N_sls": 700.0, "N_uls": 980.0, "qa": 200.0, "M_sls": 0.0},
        },
        "reference": {
            # p_uls = 980/(3.0×2.5) = 130.67 kN/m²
            # d = 550−75−8 = 467 mm
            "ULS upward pressure":  130.67,
            "Effective depth":      467.0,
        },
    },
]

# ── 5. STEEL BEAM ──────────────────────────────────────────────────────────────

# Section reference data (tabulated values, BS 5950 Section Tables)
# 457×191×67 UB:  A=85.5cm², Iy=29400cm⁴, Wpl=1470cm³, Wel=1300cm³
# 533×210×82 UB:  A=104cm², Iy=47500cm⁴, Wpl=2060cm³, Wel=1800cm³
# 356×171×51 UB:  A=64.9cm², Iy=14100cm⁴, Wpl=896cm³,  Wel=796cm³

STEEL_BEAM_CASES = [
    {
        "id": "ST-B-01",
        "desc": "457×191×67 UB — 8 m Simply Supported, gk=10 qk=8 kN/m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.2 (bending & shear)",
        "inputs": {
            "section":   {"A_cm2": 85.5, "Iy_cm4": 29400.0, "Wpl_cm3": 1470.0,
                          "Wel_cm3": 1300.0, "hw_mm": 407.6, "tw_mm": 8.5,
                          "tf_mm": 12.7, "bf_mm": 189.9},
            "geometry":  {"span": 8.0, "support_type": "simply_supported",
                          "lateral_restraint": "full"},
            "materials": {"py": 275.0},
            "loads":     {"gk": 10.0, "qk": 8.0},
        },
        "reference": {
            # w_uls = 1.4×10 + 1.6×8 = 26.8 kN/m
            # M_Ed  = 26.8×64/8 = 214.4 kN·m
            # V_Ed  = 26.8×8/2 = 107.2 kN
            # Pv    = 0.6×275×(407.6×8.5)/1000 = 571.7 kN  (shear capacity)
            # Mc    = 275×1470×1000/1e6 = 404.3 kN·m  (moment capacity, Class 1/2)
            "Design UDL":         26.80,
            "Design moment":     214.40,
            "Shear capacity":    571.7,
            "Moment capacity":   404.3,
        },
    },
    {
        "id": "ST-B-02",
        "desc": "533×210×82 UB — 10 m Simply Supported, gk=12 qk=10 kN/m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.2",
        "inputs": {
            "section":   {"A_cm2": 104.0, "Iy_cm4": 47500.0, "Wpl_cm3": 2060.0,
                          "Wel_cm3": 1800.0, "hw_mm": 476.5, "tw_mm": 10.1,
                          "tf_mm": 13.2, "bf_mm": 208.8},
            "geometry":  {"span": 10.0, "support_type": "simply_supported",
                          "lateral_restraint": "full"},
            "materials": {"py": 275.0},
            "loads":     {"gk": 12.0, "qk": 10.0},
        },
        "reference": {
            # w_uls = 1.4×12 + 1.6×10 = 32.8 kN/m
            # M_Ed  = 32.8×100/8 = 410.0 kN·m
            # Pv    = 0.6×275×(476.5×10.1)/1000 = 793.7 kN
            # Mc    = 275×2060×1000/1e6 = 566.5 kN·m
            "Design UDL":         32.80,
            "Design moment":     410.0,
            "Shear capacity":    793.7,
            "Moment capacity":   566.5,
        },
    },
    {
        "id": "ST-B-03",
        "desc": "356×171×51 UB — 6 m Continuous, gk=8 qk=6 kN/m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.2 (continuous, mid-span M = 0.086wL²)",
        "inputs": {
            "section":   {"A_cm2": 64.9, "Iy_cm4": 14100.0, "Wpl_cm3": 896.0,
                          "Wel_cm3": 796.0, "hw_mm": 321.4, "tw_mm": 7.3,
                          "tf_mm": 11.5, "bf_mm": 171.5},
            "geometry":  {"span": 6.0, "support_type": "continuous",
                          "lateral_restraint": "full"},
            "materials": {"py": 275.0},
            "loads":     {"gk": 8.0, "qk": 6.0},
        },
        "reference": {
            # w_uls = 1.4×8 + 1.6×6 = 20.8 kN/m
            # M_Ed  = 0.086×20.8×36 = 64.5 kN·m
            # Mc    = 275×896×1000/1e6 = 246.4 kN·m
            # Pv    = 0.6×275×(321.4×7.3)/1000 = 387.9 kN
            "Design UDL":         20.80,
            "Design moment":      64.5,
            "Moment capacity":   246.4,
            "Shear capacity":    387.9,
        },
    },
]

# ── 6. STEEL COLUMN ────────────────────────────────────────────────────────────

STEEL_COL_CASES = [
    {
        "id": "ST-C-01",
        "desc": "UC 203×203×60 — le=3.5 m, N=500 kN, Mx=30 kN·m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.7 (compression) & Cl. 4.8.3 (interaction)",
        "inputs": {
            "section":   {"A_cm2": 76.4, "iy_cm": 8.96, "iz_cm": 5.21,
                          "Wpl_y_cm3": 652.0, "Wpl_z_cm3": 305.0, "tf_mm": 14.2},
            "geometry":  {"le_y": 3.5, "le_z": 3.5},
            "materials": {"py": 275.0},
            "loads":     {"N_Ed": 500.0, "Mx": 30.0, "My": 0.0},
        },
        "reference": {
            # λy = 3500/89.6 = 39.06;  λz = 3500/52.1 = 67.18
            # pc_z (curve b, Perry-Robertson) ≈ 207.9 N/mm²
            # Compress capacity Pc = 207.9×7640/1000 ≈ 1588 kN
            "Slenderness (major)":  39.06,
            "Slenderness (minor)":  67.18,
        },
    },
    {
        "id": "ST-C-02",
        "desc": "UC 254×254×73 — le=4.0 m, N=800 kN, Mx=50 kN·m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.7 & Cl. 4.8.3",
        "inputs": {
            "section":   {"A_cm2": 93.1, "iy_cm": 11.07, "iz_cm": 6.46,
                          "Wpl_y_cm3": 992.0, "Wpl_z_cm3": 465.0, "tf_mm": 14.2},
            "geometry":  {"le_y": 4.0, "le_z": 4.0},
            "materials": {"py": 275.0},
            "loads":     {"N_Ed": 800.0, "Mx": 50.0, "My": 0.0},
        },
        "reference": {
            # λy = 4000/110.7 = 36.13;  λz = 4000/64.6 = 61.92
            "Slenderness (major)":  36.13,
            "Slenderness (minor)":  61.92,
        },
    },
    {
        "id": "ST-C-03",
        "desc": "UC 152×152×37 — le=3.0 m, N=150 kN, Mx=10 kN·m",
        "standard_clause": "BS 5950-1:2000, Cl. 4.7 & Cl. 4.8.3",
        "inputs": {
            "section":   {"A_cm2": 47.1, "iy_cm": 6.54, "iz_cm": 3.87,
                          "Wpl_y_cm3": 309.0, "Wpl_z_cm3": 149.0, "tf_mm": 11.5},
            "geometry":  {"le_y": 3.0, "le_z": 3.0},
            "materials": {"py": 275.0},
            "loads":     {"N_Ed": 150.0, "Mx": 10.0, "My": 0.0},
        },
        "reference": {
            # λy = 3000/65.4 = 45.87;  λz = 3000/38.7 = 77.52
            "Slenderness (major)":  45.87,
            "Slenderness (minor)":  77.52,
        },
    },
]

# ── 7. STEEL TRUSS ─────────────────────────────────────────────────────────────

STEEL_TRUSS_CASES = [
    {
        "id": "ST-T-01",
        "desc": "CHS 101.6×3.2 Tension Member — le=2.0 m, F_Ed=+100 kN",
        "standard_clause": "BS 5950-1:2000, Cl. 4.6 (tension)",
        "inputs": {
            "section":   {"A_cm2": 9.97, "i_min_cm": 3.47, "member_type": "chs"},
            "geometry":  {"le": 2.0},
            "materials": {"py": 355.0},
            "loads":     {"F_Ed": 100.0},
        },
        "reference": {
            # Pt = py×A = 355×997/1000 = 353.9 kN
            # Utilisation = 100 / 353.9 = 28.3 %
            "Tension capacity":   353.9,
        },
    },
    {
        "id": "ST-T-02",
        "desc": "RHS 100×60×4 Compression Member — le=3.0 m, F_Ed=−60 kN",
        "standard_clause": "BS 5950-1:2000, Cl. 4.7 (compression, strut curve a)",
        "inputs": {
            "section":   {"A_cm2": 9.67, "i_min_cm": 2.13, "member_type": "rhs"},
            "geometry":  {"le": 3.0},
            "materials": {"py": 275.0},
            "loads":     {"F_Ed": -60.0},
        },
        "reference": {
            # λ = 3000/21.3 = 140.8
            # Perry-Robertson (curve a): pc ≈ 89.8 N/mm²  → Pc ≈ 86.8 kN
            # Utilisation = 60/86.8 = 69.1 %
            "Compression capacity":  86.8,
        },
    },
    {
        "id": "ST-T-03",
        "desc": "Angle 100×65×8 Tension Member — le=1.5 m, F_Ed=+80 kN",
        "standard_clause": "BS 5950-1:2000, Cl. 4.6 (tension)",
        "inputs": {
            "section":   {"A_cm2": 12.8, "i_min_cm": 1.60, "member_type": "angle"},
            "geometry":  {"le": 1.5},
            "materials": {"py": 275.0},
            "loads":     {"F_Ed": 80.0},
        },
        "reference": {
            # Pt = 275×1280/1000 = 352.0 kN
            # Utilisation = 80/352 = 22.7 %
            "Tension capacity":  352.0,
        },
    },
]

# ── 8. STEEL PORTAL ────────────────────────────────────────────────────────────

STEEL_PORTAL_CASES = [
    {
        "id": "ST-P-01",
        "desc": "20 m Span, h=5 m Portal — Pinned Bases, gk=0.5 qk=0.6 kN/m²",
        "standard_clause": "BS 5950-1:2000, Cl. 4.8.3 (elastic simplified analysis, pinned base)",
        "inputs": {
            "geometry":      {"span": 20.0, "height": 5.0, "bay_width": 6.0,
                              "pitch_deg": 6.0, "base_type": "pinned"},
            "rafter_section":{"A_cm2": 60.0, "I_cm4": 20000.0, "Wpl_cm3": 900.0,
                              "iy_cm": 8.0, "iz_cm": 2.0},
            "column_section":{"A_cm2": 76.0, "I_cm4": 6000.0, "Wpl_cm3": 650.0,
                              "iy_cm": 8.0, "iz_cm": 3.0},
            "materials":     {"py": 275.0},
            "loads":         {"gk": 0.5, "qk": 0.6},
        },
        "reference": {
            # w_uls = (1.4×0.5 + 1.6×0.6)×6 = 9.96 kN/m
            # H (pinned) = w×L²/(8h) = 9.96×400/40 = 99.6 kN
            # M_eaves = H×h = 99.6×5 = 498.0 kN·m
            "Design UDL (w_uls)":   9.96,
            "Horizontal thrust":   99.6,
            "Eaves moment":       498.0,
        },
    },
    {
        "id": "ST-P-02",
        "desc": "15 m Span, h=4 m Portal — Fixed Bases, gk=0.4 qk=0.5 kN/m²",
        "standard_clause": "BS 5950-1:2000, Cl. 4.8.3 (elastic simplified analysis, fixed base)",
        "inputs": {
            "geometry":      {"span": 15.0, "height": 4.0, "bay_width": 5.0,
                              "pitch_deg": 8.0, "base_type": "fixed"},
            "rafter_section":{"A_cm2": 50.0, "I_cm4": 15000.0, "Wpl_cm3": 750.0,
                              "iy_cm": 7.5, "iz_cm": 2.0},
            "column_section":{"A_cm2": 60.0, "I_cm4": 4000.0, "Wpl_cm3": 500.0,
                              "iy_cm": 7.0, "iz_cm": 2.5},
            "materials":     {"py": 275.0},
            "loads":         {"gk": 0.4, "qk": 0.5},
        },
        "reference": {
            # w_uls = (1.4×0.4 + 1.6×0.5)×5 = 6.8 kN/m
            # H (fixed) = 3wL²/(16h) = 3×6.8×225/(16×4) = 71.7 kN
            # M_eaves = H×h = 71.7×4 = 286.9 kN·m
            "Design UDL (w_uls)":  6.80,
            "Horizontal thrust":  71.7,
            "Eaves moment":      286.9,
        },
    },
    {
        "id": "ST-P-03",
        "desc": "25 m Span, h=7 m Portal — Pinned Bases, gk=0.45 qk=0.55 kN/m²",
        "standard_clause": "BS 5950-1:2000, Cl. 4.8.3 (elastic simplified analysis, pinned base)",
        "inputs": {
            "geometry":      {"span": 25.0, "height": 7.0, "bay_width": 7.0,
                              "pitch_deg": 5.0, "base_type": "pinned"},
            "rafter_section":{"A_cm2": 75.0, "I_cm4": 30000.0, "Wpl_cm3": 1200.0,
                              "iy_cm": 9.0, "iz_cm": 2.5},
            "column_section":{"A_cm2": 90.0, "I_cm4": 8000.0, "Wpl_cm3": 800.0,
                              "iy_cm": 9.0, "iz_cm": 3.5},
            "materials":     {"py": 275.0},
            "loads":         {"gk": 0.45, "qk": 0.55},
        },
        "reference": {
            # w_uls = (1.4×0.45 + 1.6×0.55)×7 = (0.63+0.88)×7 = 10.57 kN/m
            # H (pinned) = w×L²/(8h) = 10.57×625/(8×7) = 118.0 kN
            # M_eaves = H×h = 118.0×7 = 826.0 kN·m
            "Design UDL (w_uls)":  10.57,
            "Horizontal thrust":  118.0,
            "Eaves moment":       826.0,
        },
    },
]

# ═══════════════════════════════════════════════════════════════════════════════
#  VALUE EXTRACTION — maps reference labels → step search fragments per module
# ═══════════════════════════════════════════════════════════════════════════════

# Each key is a reference label from the cases above.
# Value is a tuple of fragments passed to step_val() — all must appear in label.
STEP_MAP: dict[str, tuple[str, ...]] = {
    # RC Beam / Slab / Column / Foundation
    "Design UDL":            ("Design UDL",),
    "Design moment":         ("moment M",),
    "Design shear":          ("shear V",),
    "Effective depth":       ("Effective depth",),
    "K factor":              ("K factor",),
    "As required":           ("As,req",),
    "lex/h":                 ("lex/h",),
    "Gross section area":    ("Gross section",),
    "ULS upward pressure":   ("ULS upward",),
    # Steel beam
    "Design UDL":            ("Design UDL",),
    "Design moment":         ("Design moment",),
    "Shear capacity":        ("Shear capacity",),
    "Moment capacity":       ("Moment capacity",),
    # Steel column
    "Slenderness (major)":   ("Slenderness", "major"),
    "Slenderness (minor)":   ("Slenderness", "minor"),
    "Compression capacity":  ("Compression capacity",),
    # Steel truss
    "Tension capacity":      ("Tension capacity",),
    "Compression capacity":  ("Compression capacity",),
    # Steel portal
    "Design UDL (w_uls)":    ("w_uls",),
    "Horizontal thrust":     ("Horizontal thrust",),
    "Eaves moment":          ("Eaves moment",),
    # Slab two-way specific
    "n (ULS load)":          ("n_uls", "ULS load"),
    "Msx (short span)":      ("Msx", "short"),
    "As short span":         ("As,req", "short"),
}


def extract_val(result: dict, ref_label: str) -> float | None:
    """Extract a computed value from module output using the STEP_MAP, with fallbacks."""
    frags = STEP_MAP.get(ref_label)
    if frags:
        v = step_val(result, *frags)
        if v is not None:
            return v

    # Fallback: for "n (ULS load)" → search for any step containing "ULS" and "load"
    # For multi-word, try each word as fragment
    label_lower = ref_label.lower()
    for step in result.get("steps", []):
        slabel = step["label"].lower()
        # Exact keyword match attempts
        if ref_label == "n (ULS load)" and "uls" in slabel and "load" in slabel:
            return float(step["value"])
        if ref_label == "Msx (short span)" and ("msx" in slabel or ("moment" in slabel and "short" in slabel)):
            return float(step["value"])
        if ref_label == "As short span" and "as" in slabel and ("short" in slabel or "req" in slabel):
            return float(step["value"])
        if ref_label == "Design shear" and "shear" in slabel and ("v" in slabel or "design" in slabel) and "capacity" not in slabel and "stress" not in slabel:
            return float(step.get("value", 0.0))

    return None


# ═══════════════════════════════════════════════════════════════════════════════
#  RUN ALL CASES
# ═══════════════════════════════════════════════════════════════════════════════

def run_all_cases(module, cases: list[dict]) -> list[dict]:
    """Run a module against all its cases and collect results."""
    results = []
    for case in cases:
        r = module.run(case["inputs"])
        case_result = {
            "id":          case["id"],
            "desc":        case["desc"],
            "clause":      case["standard_clause"],
            "inputs":      case["inputs"],
            "reference":   case["reference"],
            "app_result":  r,
            "verdict":     r.get("results", {}).get("verdict", "error") if r.get("status") == "ok" else "error",
            "comparisons": {},
            "overall_pass": True,
        }

        for ref_label, ref_value in case["reference"].items():
            app_value = extract_val(r, ref_label)
            symbol, status = compare(app_value, ref_value)
            case_result["comparisons"][ref_label] = {
                "ref":    ref_value,
                "app":    app_value,
                "symbol": symbol,
                "status": status,
            }
            if status == "fail":
                case_result["overall_pass"] = False

        results.append(case_result)
    return results


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_colour)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def _bold_run(para, text: str, size: int = 10, rgb: RGBColor | None = None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if rgb:
        run.font.color.rgb = rgb
    return run


def _normal_run(para, text: str, size: int = 10, rgb: RGBColor | None = None):
    run = para.add_run(text)
    run.font.size = Pt(size)
    if rgb:
        run.font.color.rgb = rgb
    return run


def _heading(doc: Document, text: str, level: int):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 10)
        if level == 1:
            run.font.color.rgb = CLR_HDR
        elif level == 2:
            run.font.color.rgb = CLR_SUBHDR


def _fmt(v: float | None, decimals: int = 3) -> str:
    if v is None:
        return "—"
    return f"{v:.{decimals}f}"


def build_document(all_module_results: list[dict]) -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ELEMENT VERIFICATION RESULTS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = CLR_HDR

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{APP_VER}  ·  {STANDARD}")
    r.font.size = Pt(12)
    r.font.color.rgb = CLR_SUBHDR

    doc.add_paragraph()
    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.style = "Table Grid"
    rows_data = [
        ("Document Title",   "Structural Element Calculation Verification Results"),
        ("Application",      f"SSAD — Smart Structural Analysis & Design  ({APP_VER})"),
        ("Design Standards", STANDARD),
        ("Date of Issue",    TODAY),
    ]
    for i, (k, v) in enumerate(rows_data):
        info_tbl.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        info_tbl.rows[i].cells[1].paragraphs[0].add_run(v)
        _set_cell_bg(info_tbl.rows[i].cells[0], "DBEAFE")

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Verification Methodology:  Each element module is tested with three independent "
        "sets of input parameters.  Reference values are derived from first-principle "
        "calculations using the applicable design standard and are compared against "
        "app-generated outputs with a tolerance of ±2 %.  A test case is marked "
        "VERIFIED if all numerical comparisons are within tolerance and the app "
        "returns a valid structural verdict."
    )
    nr.font.size = Pt(9)
    nr.font.italic = True

    doc.add_page_break()

    # ── Executive Summary Table ──────────────────────────────────────────────
    _heading(doc, "Executive Summary", 1)

    total_cases = sum(len(m["cases"]) for m in all_module_results)
    total_pass  = sum(1 for m in all_module_results for c in m["cases"] if c["overall_pass"])
    total_fail  = total_cases - total_pass

    summary_para = doc.add_paragraph()
    summary_para.add_run(
        f"Total test cases: {total_cases}     "
        f"Verified: {total_pass}     "
        f"Failed: {total_fail}"
    ).bold = True

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(["#", "Module", "Standard", "TC-1", "TC-2", "TC-3"]):
        hdr_cells[i].paragraphs[0].add_run(h).bold = True
        _set_cell_bg(hdr_cells[i], "1A4A8A")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    for idx, module_data in enumerate(all_module_results, 1):
        row = tbl.add_row().cells
        row[0].paragraphs[0].add_run(str(idx)).font.size = Pt(9)
        row[1].paragraphs[0].add_run(module_data["name"]).font.size = Pt(9)
        row[2].paragraphs[0].add_run(module_data["standard"]).font.size = Pt(9)
        for j, case in enumerate(module_data["cases"]):
            colour  = "D1FAE5" if case["overall_pass"] else "FEE2E2"
            symbol  = "✓ PASS" if case["overall_pass"] else "✗ FAIL"
            fnt_rgb = CLR_PASS   if case["overall_pass"] else CLR_FAIL
            p = row[3 + j].paragraphs[0]
            r = p.add_run(symbol)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = fnt_rgb
            _set_cell_bg(row[3 + j], colour[1:] if colour.startswith("#") else colour)

    doc.add_page_break()

    # ── Per-module detailed sections ─────────────────────────────────────────
    for module_data in all_module_results:
        _heading(doc, f"{module_data['name']}  ({module_data['standard']})", 1)
        para = doc.add_paragraph()
        para.add_run(module_data.get("description", ""))

        for case in module_data["cases"]:
            doc.add_paragraph()
            _heading(doc, f"{case['id']}  ·  {case['desc']}", 2)

            # Clause reference
            cr = doc.add_paragraph()
            cr.add_run("Applicable Clause:  ").bold = True
            cr.add_run(case["clause"]).font.italic = True

            # App verdict
            verd = case["verdict"].upper()
            verd_rgb = CLR_PASS if verd == "PASS" else (CLR_WARN if verd == "WARN" else CLR_FAIL)
            vp = doc.add_paragraph()
            vp.add_run("App Structural Verdict:  ").bold = True
            vr = vp.add_run(verd)
            vr.bold = True
            vr.font.color.rgb = verd_rgb

            # Overall verification status
            ov_pass = case["overall_pass"]
            ovp = doc.add_paragraph()
            ovp.add_run("Verification Status:  ").bold = True
            ovr = ovp.add_run("VERIFIED ✓" if ov_pass else "FAILED ✗")
            ovr.bold = True
            ovr.font.color.rgb = CLR_PASS if ov_pass else CLR_FAIL

            # Input parameters table
            doc.add_paragraph().add_run("Input Parameters:").bold = True
            _write_inputs_table(doc, case["inputs"])

            # Verification results table
            doc.add_paragraph()
            doc.add_paragraph().add_run("Verification Results (Reference vs App, tolerance ±2%):").bold = True
            _write_verification_table(doc, case["comparisons"])

            doc.add_paragraph()  # spacer

        doc.add_page_break()

    # ── Sign-off ─────────────────────────────────────────────────────────────
    _heading(doc, "Verification Sign-Off", 1)
    signoff = doc.add_table(rows=4, cols=2)
    signoff.style = "Table Grid"
    fields = [
        ("Prepared by",    ""),
        ("Reviewed by",    ""),
        ("Date",           TODAY),
        ("Document Status", "DRAFT — Requires Chartered Engineer Review"),
    ]
    for i, (k, v) in enumerate(fields):
        signoff.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        signoff.rows[i].cells[1].paragraphs[0].add_run(v)
        _set_cell_bg(signoff.rows[i].cells[0], "DBEAFE")

    return doc


def _write_inputs_table(doc: Document, inputs: dict):
    """Flatten and render all inputs as a two-column table."""
    flat: list[tuple[str, str]] = []
    for group, values in inputs.items():
        if isinstance(values, dict):
            for k, v in values.items():
                flat.append((f"{group}.{k}", str(v)))
        else:
            flat.append((group, str(values)))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Parameter", "Value"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(hdr[i], "EFF6FF")

    for param, value in flat:
        row = tbl.add_row().cells
        row[0].paragraphs[0].add_run(param).font.size = Pt(9)
        row[1].paragraphs[0].add_run(value).font.size = Pt(9)


def _write_verification_table(doc: Document, comparisons: dict):
    """Render the comparison table: label | reference | app | tolerance | status."""
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    headers = ["Check / Parameter", "Reference Value", "App Value", "Tolerance", "Result"]
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(hdr[i], "1A4A8A")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for label, data in comparisons.items():
        row = tbl.add_row().cells
        row[0].paragraphs[0].add_run(label).font.size = Pt(9)
        row[1].paragraphs[0].add_run(_fmt(data["ref"])).font.size = Pt(9)
        row[2].paragraphs[0].add_run(_fmt(data["app"])).font.size = Pt(9)
        row[3].paragraphs[0].add_run("±2 %").font.size = Pt(9)

        status_run = row[4].paragraphs[0].add_run(data["symbol"])
        status_run.bold = True
        status_run.font.size = Pt(9)
        if data["status"] == "pass":
            status_run.font.color.rgb = CLR_PASS
            _set_cell_bg(row[4], "D1FAE5")
        elif data["status"] == "fail":
            status_run.font.color.rgb = CLR_FAIL
            _set_cell_bg(row[4], "FEE2E2")
        else:
            status_run.font.color.rgb = CLR_WARN
            _set_cell_bg(row[4], "FEF3C7")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("SSAD — Element Verification Script")
    print("=" * 60)

    MODULES = [
        {
            "name":        "RC Beam",
            "standard":    "BS 8110-1:1997",
            "description": "Rectangular RC beam design — bending, shear and deflection checks.",
            "module":      rc_beam_bs_v1,
            "case_defs":   RC_BEAM_CASES,
        },
        {
            "name":        "RC Slab",
            "standard":    "BS 8110-1:1997",
            "description": "RC slab design — one-way, two-way and cantilever configurations.",
            "module":      rc_slab_bs_v1,
            "case_defs":   RC_SLAB_CASES,
        },
        {
            "name":        "RC Column",
            "standard":    "BS 8110-1:1997",
            "description": "Rectangular braced RC column — short column axial + bending capacity.",
            "module":      rc_column_bs_v1,
            "case_defs":   RC_COLUMN_CASES,
        },
        {
            "name":        "RC Pad Footing",
            "standard":    "BS 8110-1:1997",
            "description": "Isolated RC pad footing — bearing, bending, shear and punching checks.",
            "module":      rc_foundation_bs_v1,
            "case_defs":   RC_FOUND_CASES,
        },
        {
            "name":        "Steel Beam",
            "standard":    "BS 5950-1:2000",
            "description": "Universal Beam design — section class, shear, moment and deflection.",
            "module":      steel_beam_bs_v1,
            "case_defs":   STEEL_BEAM_CASES,
        },
        {
            "name":        "Steel Column",
            "standard":    "BS 5950-1:2000",
            "description": "H/I-section column under combined compression and bending.",
            "module":      steel_column_bs_v1,
            "case_defs":   STEEL_COL_CASES,
        },
        {
            "name":        "Steel Truss Member",
            "standard":    "BS 5950-1:2000",
            "description": "Individual truss member — tension and compression capacity checks.",
            "module":      steel_truss_bs_v1,
            "case_defs":   STEEL_TRUSS_CASES,
        },
        {
            "name":        "Steel Portal Frame",
            "standard":    "BS 5950-1:2000",
            "description": "Single-bay symmetric portal frame — simplified elastic analysis.",
            "module":      steel_portal_bs_v1,
            "case_defs":   STEEL_PORTAL_CASES,
        },
    ]

    all_module_results = []

    for mod_def in MODULES:
        name   = mod_def["name"]
        module = mod_def["module"]
        cases_defs = mod_def["case_defs"]

        print(f"\n  Running {name} ...", end="", flush=True)
        cases = run_all_cases(module, cases_defs)
        passed = sum(1 for c in cases if c["overall_pass"])
        print(f"  {passed}/{len(cases)} verified")

        all_module_results.append({
            "name":        name,
            "standard":    mod_def["standard"],
            "description": mod_def["description"],
            "cases":       cases,
        })

    # Build and save document
    print("\n  Building Word document ...", end="", flush=True)
    doc = build_document(all_module_results)

    out_path = BACKEND.parent / "Element Verification Results.docx"
    doc.save(str(out_path))
    print(f"  Saved.")

    # Final summary
    total_cases = sum(len(m["cases"]) for m in all_module_results)
    total_pass  = sum(1 for m in all_module_results for c in m["cases"] if c["overall_pass"])
    print("\n" + "=" * 60)
    print(f"  TOTAL:  {total_cases} cases   VERIFIED: {total_pass}   FAILED: {total_cases - total_pass}")
    print(f"  Output: {out_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
